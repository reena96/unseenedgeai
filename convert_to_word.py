#!/usr/bin/env python3
"""Convert MASS_TECHNICAL_OVERVIEW.md to Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document with formatting."""

    # Read markdown file
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split into lines
    lines = content.split("\n")

    in_table = False
    table_headers = []
    table_data = []
    in_code_block = False
    code_lines = []

    for line in lines:
        # Handle code blocks
        if line.startswith("```"):
            if in_code_block:
                # End of code block
                if code_lines:
                    code_para = doc.add_paragraph("\n".join(code_lines))
                    code_para.style = "List Bullet"
                    code_para_format = code_para.paragraph_format
                    code_para_format.left_indent = Inches(0.5)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle tables
        if line.startswith("|") and not line.startswith("|---"):
            if not in_table:
                in_table = True
                # Parse headers
                table_headers = [cell.strip() for cell in line.split("|")[1:-1]]
                table_data = []
            else:
                # Parse data row
                row_data = [cell.strip() for cell in line.split("|")[1:-1]]
                table_data.append(row_data)
            continue
        elif line.startswith("|---"):
            continue
        elif in_table and not line.startswith("|"):
            # End of table - create it
            if table_headers and table_data:
                table = doc.add_table(rows=len(table_data) + 1, cols=len(table_headers))
                table.style = "Light Grid Accent 1"

                # Add headers
                for i, header in enumerate(table_headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True

                # Add data
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx + 1].cells[col_idx].text = cell_data

                doc.add_paragraph()  # Add spacing after table

            in_table = False
            table_headers = []
            table_data = []

        # Skip empty lines in tables
        if in_table:
            continue

        # Handle headings
        if line.startswith("# ") and not line.startswith("## "):
            # H1
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            # H2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            # H3
            doc.add_heading(line[4:], level=3)

        # Handle horizontal rules
        elif line.startswith("---"):
            doc.add_paragraph("_" * 80)

        # Handle bullet lists
        elif line.startswith("- **") or line.startswith("• **"):
            # Bold bullet point
            text = line.lstrip("- •").strip()
            p = doc.add_paragraph(style="List Bullet")
            # Parse bold text
            parts = re.split(r"\*\*(.*?)\*\*", text)
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are bold
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)

        elif line.startswith("- ") or line.startswith("• "):
            # Regular bullet point
            doc.add_paragraph(line.lstrip("- •").strip(), style="List Bullet")

        # Handle regular paragraphs
        elif line.strip() and not line.startswith("#"):
            # Parse inline formatting
            p = doc.add_paragraph()

            # Handle bold text
            parts = re.split(r"\*\*(.*?)\*\*", line)
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are bold
                    run = p.add_run(part)
                    run.bold = True
                else:
                    # Handle inline code
                    code_parts = re.split(r"`(.*?)`", part)
                    for j, code_part in enumerate(code_parts):
                        if j % 2 == 1:  # Inline code
                            run = p.add_run(code_part)
                            run.font.name = "Courier New"
                            run.font.size = Pt(9)
                        else:
                            p.add_run(code_part)

    # Save document
    doc.save(docx_file)
    print(f"✅ Converted {md_file} to {docx_file}")


if __name__ == "__main__":
    parse_markdown_to_docx("MASS_TECHNICAL_OVERVIEW.md", "MASS_TECHNICAL_OVERVIEW.docx")
